"""Genera el documento del Proyecto Final (.docx) con EDA, figuras y espacios para capturas."""
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

REPO = Path("..")
FIG = REPO / "figures"
ASSETS = Path("doc_assets")
stats = json.loads((ASSETS / "eda_stats.json").read_text(encoding="utf-8"))
OUT = REPO.resolve().parent / "Proyecto_Final_ECG.docx"

URL = "https://JoeUgalde40-ecg-cnn1d-api.hf.space"
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()
# márgenes
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h(text, size=15, color=ACCENT, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def para(text, size=11, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def figure(path, caption, width=6.0):
    if not Path(path).exists():
        para(f"[falta figura: {path}]", color=GRAY, italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(caption, size=9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(10)


def placeholder(num, titulo, instruccion, alto_cm=7.5):
    """Inserta la captura cap{num}.png si existe; si no, una caja gris rotulada."""
    cap_path = ASSETS / f"cap{num}.png"
    if cap_path.exists():
        import struct
        with open(cap_path, "rb") as f:
            f.read(16)
            w, hgt = struct.unpack(">II", f.read(8))
        # imágenes muy verticales (swagger) más angostas para que quepan en la página
        width = 4.0 if hgt > w * 1.4 else 6.2
        figure(cap_path, f"Figura. {titulo}", width=width)
        return
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    cell.width = Inches(6.2)
    shade(cell, "EFEFEF")
    # borde
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed"); e.set(qn("w:sz"), "8")
        e.set(qn("w:color"), "AAAAAA")
        borders.append(e)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"\n📷  CAPTURA {num}\n")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT
    r2 = p.add_run(f"{instruccion}\n\n(pega aquí tu captura de pantalla)\n")
    r2.font.size = Pt(10); r2.font.color.rgb = GRAY
    cap = para(f"Figura. {titulo}", size=9, italic=True, color=GRAY,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(12)


# ===================== PORTADA =====================
sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(70)
t = para("Clasificación de Electrocardiogramas mediante\nRedes Convolucionales 1-D",
         size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
t.runs[0].bold = True; t.runs[0].font.color.rgb = ACCENT
para("Una aplicación de clasificación vía API", size=14, italic=True,
     color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_before = Pt(40)
para("Proyecto Final", size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Aprendizaje Profundo — Maestría en Ciencia de Datos",
     size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_before = Pt(30)
para("Integrantes: ___________________________________________",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Fecha: junio de 2026", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_before = Pt(20)
lk = para(f"Aplicación en línea: {URL}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
lk.runs[0].font.color.rgb = RGBColor(0x21, 0x5F, 0x9A)
doc.add_page_break()

# ===================== 1. INTRODUCCIÓN =====================
h("1. Introducción y planteamiento del problema")
para(
    "El electrocardiograma (ECG) es la prueba más común para evaluar la actividad eléctrica del "
    "corazón y detectar patologías cardíacas. Su interpretación, sin embargo, requiere personal "
    "especializado y es susceptible a variabilidad entre observadores. Este proyecto desarrolla un "
    "sistema de aprendizaje profundo que clasifica automáticamente ECGs de 12 derivaciones en cinco "
    "categorías diagnósticas, y lo expone como una aplicación accesible vía API.")
para(
    "El problema es de clasificación multi-etiqueta: un mismo ECG puede presentar varias condiciones "
    "simultáneamente. El modelo entrega, para cada ECG, una probabilidad independiente por cada una de "
    "las cinco superclases diagnósticas.")

# ===================== 2. DATASET =====================
h("2. Conjunto de datos: PTB-XL")
para(
    "Se utilizó PTB-XL, una base de datos pública de PhysioNet ampliamente usada en investigación. "
    "Contiene registros de ECG clínicos de 12 derivaciones, cada uno de 10 segundos de duración, "
    "muestreados a 100 Hz (y también a 500 Hz). Cada registro está anotado por cardiólogos y agrupado "
    "en cinco superclases diagnósticas.")
para("Las cinco superclases diagnósticas son:")
tbl = doc.add_table(rows=6, cols=2); tbl.style = "Light Grid Accent 1"
tbl.cell(0, 0).paragraphs[0].add_run("Superclase").bold = True
tbl.cell(0, 1).paragraphs[0].add_run("Significado clínico").bold = True
filas = [("NORM", "ECG normal (sano)"), ("MI", "Infarto de miocardio"),
         ("STTC", "Alteraciones del segmento ST / onda T"),
         ("CD", "Trastornos de la conducción eléctrica"),
         ("HYP", "Hipertrofia (crecimiento anormal del músculo cardíaco)")]
for i, (c, d) in enumerate(filas, 1):
    tbl.cell(i, 0).text = c
    tbl.cell(i, 1).text = d
doc.add_paragraph()
figure(FIG / "ecg_paciente_12_derivaciones.png",
       "Ejemplo de una entrada del modelo: ECG real de 12 derivaciones (paciente sano, clase NORM). "
       "A la izquierda las derivaciones de miembros; a la derecha las precordiales.", width=6.2)

# ===================== 3. EDA =====================
doc.add_page_break()
h("3. Análisis exploratorio de datos (EDA)")
para(
    f"El conjunto contiene {stats['n_ecgs']:,} ECGs provenientes de {stats['n_pacientes']:,} pacientes. "
    f"La población está equilibrada por sexo ({stats['pct_hombres']:.0f}% hombres, "
    f"{stats['pct_mujeres']:.0f}% mujeres), con una edad mediana de {stats['edad_mediana']:.0f} años. "
    "Para evitar fuga de información, la partición train/validación/test se hizo por paciente: ningún "
    "paciente aparece en más de un conjunto.")

h("3.1 Distribución de las clases", size=12, color=GRAY, space_before=8)
para(
    "Las clases están notablemente desbalanceadas. NORM es la mayoritaria "
    f"({stats['pct_clases']['NORM']}%), mientras que HYP es la más escasa "
    f"({stats['pct_clases']['HYP']}%). Este desbalance obliga a usar pesos de clase durante el "
    "entrenamiento y a evaluar con métricas macro (F1-macro, AUC-macro) en lugar de exactitud simple.")
figure(ASSETS / "eda_clases.png",
       "Frecuencia de cada superclase diagnóstica en el conjunto PTB-XL.", width=5.6)

h("3.2 Naturaleza multi-etiqueta", size=12, color=GRAY, space_before=8)
me = stats["multietiqueta"]
para(
    f"La mayoría de los ECGs tiene una sola etiqueta ({me['1_etiqueta']:,}), pero "
    f"{me['2_etiquetas']:,} tienen dos y {me['3+_etiquetas']:,} tienen tres o más superclases "
    "simultáneas. Esto confirma que el problema debe abordarse como clasificación multi-etiqueta y no "
    "como clasificación de clase única.")
figure(ASSETS / "eda_multietiqueta.png",
       "Número de superclases por ECG: la coexistencia de patologías es frecuente.", width=4.8)

h("3.3 Distribución de edad y sexo", size=12, color=GRAY, space_before=8)
para(
    "La distribución de edad se concentra en adultos mayores, coherente con una población clínica. "
    "Ambos sexos están bien representados a lo largo de todo el rango de edad.")
figure(ASSETS / "eda_edad.png",
       "Distribución de edad por sexo. La línea punteada marca la mediana.", width=5.6)

# ===================== 4. JUSTIFICACIÓN =====================
doc.add_page_break()
h("4. Justificación de la herramienta")
para(
    "El ECG es una señal temporal multicanal (12 derivaciones × 1000 muestras). Para este tipo de dato, "
    "las redes neuronales convolucionales 1-D (CNN 1-D) —vistas en clase— son la herramienta natural: "
    "las convoluciones a lo largo del tiempo aprenden a detectar patrones morfológicos locales del "
    "latido (pendientes, picos del complejo QRS, alteraciones del segmento ST) de forma equivalente a "
    "como las CNN 2-D detectan bordes en imágenes, pero sobre el eje temporal.")
para(
    "Frente a métodos clásicos que requieren extraer manualmente características (intervalos, amplitudes), "
    "la CNN 1-D aprende automáticamente las representaciones relevantes desde la señal cruda, lo que "
    "justifica plenamente el uso de aprendizaje profundo para este problema.")

# ===================== 5. MODELO =====================
h("5. Arquitectura del modelo")
para(
    "El clasificador combina dos ramas. La rama convolucional procesa la señal de 12 derivaciones a "
    "través de bloques Conv1D + BatchNorm + ReLU + MaxPool, y la resume en un vector de 256 "
    "características mediante global average pooling. La rama de metadatos aporta edad y sexo del "
    "paciente. Ambas se concatenan en un embedding de 258 dimensiones que pasa a una cabeza densa con "
    "salida sigmoide de 5 unidades: una probabilidad independiente por superclase.")
figure(FIG / "arquitectura_dos_ramas.png",
       "Arquitectura de dos ramas: señal (CNN 1-D) + metadatos clínicos.", width=6.0)

h("5.1 Desempeño", size=12, color=GRAY, space_before=8)
para(
    "El modelo desplegado (100 Hz) alcanza F1-macro ≈ 0.75 y AUC-macro ≈ 0.93 en el conjunto de test "
    "(con umbrales óptimos por clase calculados en validación). NORM es la clase más fácil "
    "(F1 ≈ 0.87) y HYP la más difícil, por ser minoritaria y de morfología sutil.")
tbl2 = doc.add_table(rows=4, cols=3); tbl2.style = "Light Grid Accent 1"
hdr = ["Métrica", "100 Hz", "500 Hz"]
for j, x in enumerate(hdr):
    tbl2.cell(0, j).paragraphs[0].add_run(x).bold = True
for i, (m, a, b) in enumerate([("F1-macro", "0.754", "0.767"),
                               ("F1-micro", "0.781", "0.792"),
                               ("AUC-macro", "0.932", "0.937")], 1):
    tbl2.cell(i, 0).text = m; tbl2.cell(i, 1).text = a; tbl2.cell(i, 2).text = b
doc.add_paragraph()
para(
    "Resolución 100 Hz vs. 500 Hz: la diferencia de desempeño entre entrenar con la señal a 100 Hz o a "
    "500 Hz es marginal (~1–2 % en F1-macro). Incluso para diagnósticos específicos como HYP, la mejora "
    "no justifica el ~5× de cómputo y almacenamiento que implica la mayor resolución. Por tanto, 100 Hz "
    "es la elección práctica para el despliegue, logrando cerca del 98 % del desempeño a una fracción del "
    "costo; por esta razón el modelo puesto en producción es el de 100 Hz.")

h("5.2 Interpretabilidad: ¿en qué se fija el modelo? (Grad-CAM)", size=12, color=GRAY, space_before=10)
para(
    "Para verificar que el modelo aprende patrones clínicamente razonables y no atajos espurios, se aplicó "
    "Grad-CAM 1-D sobre la última capa convolucional. La técnica produce un mapa de calor a lo largo del "
    "tiempo: las zonas en color cálido (amarillo/rojo) indican las regiones de la señal que más influyeron "
    "en la decisión del modelo para cada diagnóstico.")
para(
    "La figura siguiente muestra el Grad-CAM para el ECG #274 (etiqueta real MI, STTC, HYP; el modelo "
    "detectó las tres correctamente). Se observa que la atención se concentra en los complejos del latido "
    "—la región fisiológicamente relevante para estas patologías— y no en zonas de ruido entre latidos, "
    "lo que respalda la validez de las decisiones del modelo.")
figure(ASSETS / "gradcam_274.png",
       "Grad-CAM 1-D del ECG #274 para cada diagnóstico detectado (derivación II). "
       "Cálido = región donde el modelo más se fija para decidir.", width=6.3)

# ===================== 6. APLICACIÓN / API =====================
doc.add_page_break()
h("6. La aplicación vía API")
para(
    "El modelo entrenado se expone como una aplicación vía API construida con FastAPI (Python), "
    "empaquetada en un contenedor Docker y desplegada de forma pública y gratuita en Hugging Face "
    "Spaces. La aplicación incluye tanto una interfaz web de demostración como una API REST documentada.")
para("Está disponible en línea en:")
lk2 = para(f"   {URL}/        (interfaz web)", size=11)
lk2.runs[0].font.color.rgb = RGBColor(0x21, 0x5F, 0x9A)
lk3 = para(f"   {URL}/docs    (documentación Swagger de la API)", size=11)
lk3.runs[0].font.color.rgb = RGBColor(0x21, 0x5F, 0x9A)
para("Principales endpoints de la API:")
tbl3 = doc.add_table(rows=5, cols=3); tbl3.style = "Light Grid Accent 1"
for j, x in enumerate(["Método", "Ruta", "Descripción"]):
    tbl3.cell(0, j).paragraphs[0].add_run(x).bold = True
eps = [("GET", "/health", "Estado del servicio y modelo cargado"),
       ("GET", "/samples", "Lista de ECGs de muestra disponibles"),
       ("GET", "/plot/{id}", "Gráfica PNG de las 12 derivaciones"),
       ("POST", "/predict", "Clasifica un ECG y devuelve las probabilidades")]
for i, (a, b, c) in enumerate(eps, 1):
    tbl3.cell(i, 0).text = a; tbl3.cell(i, 1).text = b; tbl3.cell(i, 2).text = c
doc.add_paragraph()
para(
    "Flujo de uso: el usuario selecciona un ECG, la aplicación lo preprocesa (normalización por "
    "derivación e incorporación de edad y sexo), lo pasa por la red y muestra el diagnóstico junto con "
    "la probabilidad de cada superclase y la gráfica de la señal.")

# ===================== 7. CAPTURAS =====================
doc.add_page_break()
h("7. Funcionamiento de la aplicación (capturas de pantalla)")
para("A continuación se muestran capturas de la aplicación en funcionamiento.", italic=True, color=GRAY)

placeholder(1, "Pantalla principal de la interfaz web de la aplicación.",
            "Abre " + URL + "/ y captura la pantalla inicial con el selector de ECG.")
placeholder(2, "Clasificación de un ECG con múltiples patologías (caso multi-etiqueta).",
            "Elige un ECG multi-etiqueta, presiona «Clasificar» y captura el resultado: "
            "gráfica + diagnóstico + barras de probabilidad.")
placeholder(3, "Clasificación de un ECG normal (NORM) para contraste.",
            "Elige un ECG con etiqueta real NORM, clasifícalo y captura el resultado.")
placeholder(4, "Documentación interactiva (Swagger) de la API REST.",
            "Abre " + URL + "/docs y captura la vista con todos los endpoints.")
placeholder(5, "Respuesta del endpoint POST /predict ejecutado desde Swagger.",
            "En /docs, expande POST /predict → «Try it out» → envía {\"ecg_id\": 146} "
            "y captura la respuesta JSON.")
para(
    "Además, la aplicación integra la visualización Grad-CAM directamente en la interfaz: tras "
    "clasificar un ECG, el botón «Ver Grad-CAM» muestra el mapa de calor por cada clase detectada, "
    "indicando en qué regiones del latido se fija el modelo para decidir.", color=GRAY, italic=True)
placeholder(6, "Grad-CAM integrado en la aplicación web (ECG #274 con sus diagnósticos y el mapa de calor por clase).",
            "Clasifica el ECG #274 y presiona «Ver Grad-CAM»; captura el resultado completo.")

# ===================== 8. CONCLUSIONES =====================
doc.add_page_break()
h("8. Conclusiones")
para(
    "Se construyó un sistema completo de clasificación de electrocardiogramas: desde el análisis del "
    "conjunto PTB-XL y el entrenamiento de una CNN 1-D de dos ramas, hasta su despliegue como una "
    "aplicación pública vía API. El modelo resuelve un problema clínico real de clasificación "
    "multi-etiqueta con desempeño sólido (F1-macro ≈ 0.75), y la aplicación permite obtener un "
    "diagnóstico a partir de un ECG de forma inmediata y accesible desde el navegador.")
para(
    "El uso de redes convolucionales 1-D queda justificado por la naturaleza temporal y multicanal de "
    "la señal, y el despliegue mediante FastAPI + Docker en Hugging Face Spaces cumple el requisito de "
    "una aplicación funcional accesible vía API.")

doc.save(str(OUT))
print("Documento guardado en:", OUT.resolve())
